# ------------

from dotenv import load_dotenv
import os
import openai

# -------------

import magic
import docx
import json
import PyPDF2


# -------------

load_dotenv()

openai.api_key = os.getenv("OPENAI_API_KEY")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.llms import OpenAI
from langchain.chains.question_answering import load_qa_chain
from langchain.callbacks import get_openai_callback
import os
import openai

def read_file_to_string(file_path):
    try:
        # Initialize an empty string to store the file contents
        file_contents = ""

        # Open the file and read its contents
        with open(file_path, 'r', encoding='utf-8') as file:
            file_contents = file.read()

        return file_contents
    except FileNotFoundError:
        return "File not found."


def create_vector(content, vector_folder_name):

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
        )
    chunks = text_splitter.split_text(text=content)
    
    embeddings = OpenAIEmbeddings()
    VectorStore = FAISS.from_texts(chunks, embedding=embeddings)
    
    # os.makedirs(os.path.dirname(f"{vector_name}.pkl"), exist_ok=True)
    
    # directory = os.path.dirname(vector_name)
    # if not os.path.exists(directory):
    #     os.makedirs(directory)
        
    VectorStore.save_local(vector_folder_name)
    
    # with open(f"{vector_name}.pkl", "wb") as f:
    #     print("Dumping in ",f"{vector_name}.pkl")
    #     pickle.dump(VectorStore, f)

    return vector_folder_name




def query_from_vector(query, userId, filename):

    vector_folder_name = f"data/{userId}/vectors/{filename}"

    # if os.path.exists(f"{vector_name}.pkl"):
    #     with open(f"{vector_name}.pkl", "rb") as f:
    #         VectorStore = pickle.load(f)
    
    embeddings = OpenAIEmbeddings()    
    VectorStore = FAISS.load_local(vector_folder_name, embeddings=embeddings)


    # query = "what colour is the sky"

    docs = VectorStore.similarity_search(query=query, k=3)

    llm = OpenAI()
    chain = load_qa_chain(llm=llm, chain_type="stuff")
    with get_openai_callback() as cb:
        response = chain.run(input_documents=docs, question=query)
        print(cb)
    # print("\n\nResponse : ",response)
    return response


def read_document(file_path):
    # Use magic to determine the file type
    # mime = magic.Magic()
    # mime = magic.Magic(magic_file=file_path)
    file_type = magic.from_file(file_path)
    
    content = ''

    # Read content based on the file type
    if file_type.startswith('Microsoft Word'):
        # # For .docx files
        # doc = docx.Document(file_path)
        # content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        if file_path.endswith('.docx'):
            # For .docx files
            doc = docx.Document(file_path)
            content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_type.startswith('JSON'):
        # For .json files
        with open(file_path, 'r') as json_file:
            content = json.load(json_file)
    elif file_type.startswith('ASCII'):
        # For plain text (.txt) files
        with open(file_path, 'r') as txt_file:
            content = txt_file.read()
    elif file_type.startswith('PDF'):
        pdfFileObj = open(file_path, 'rb')
 
        # creating a pdf reader object
        pdfReader = PyPDF2.PdfReader(pdfFileObj)
 
        # printing number of pages in pdf file
        # print(len(pdfReader.pages))

        # creating a page object
        pageObj = pdfReader.pages[0]

        # extracting text from page
        content = pageObj.extract_text()

        # closing the pdf file object
        pdfFileObj.close()

    return content

def merge_db(user_id):
    vector_base_folder = f"data/{user_id}/vectors"
    final_folder = f"data/{user_id}/merged_vector"
    embeddings = OpenAIEmbeddings()
    all_items  = os.listdir(vector_base_folder)
    folders = [item for item in all_items if os.path.isdir(os.path.join(vector_base_folder, item))]
    if len(folders)==1:
        return "Merged - Single"
    VectorStore1 = FAISS.load_local(f"{vector_base_folder}/{folders[0]}", embeddings=embeddings)
    VectorStore2 = FAISS.load_local(f"{vector_base_folder}/{folders[1]}", embeddings=embeddings)
    VectorStore2.merge_from(VectorStore1)
    VectorStore2.save_local(final_folder)
    for i in range(1,len(folders)):
        VectorStore1 = FAISS.load_local(final_folder, embeddings=embeddings)
        VectorStore2 = FAISS.load_local(f"{vector_base_folder}/{folders[i]}", embeddings=embeddings)
        VectorStore2.merge_from(VectorStore1)
        VectorStore2.save_local(final_folder)

    response = "Merged - Multiple"

    return response


# print(query_from_vector("give me a summary about networking", r"data/uuid2502/merged"))

# vector_base_folder = "data/uuid2502/vectors"

# filename = "test"

# vector_folder_name = f"{vector_base_folder}/{filename}"

# print(create_vector(read_document(r"data/uuid2502/uploads/sample.txt"), vector_folder_name))
# print(query_from_vector("give me summary about operating system", vector_folder_name))

# print(read_document(r"data/uuid2502/uploads/docx_file.docx"))
# vector_base_folder = r"data/uuid2502/vectors"

# vector_folder_name = f"{vector_base_folder}/{filename}"

# final_folder = r"data/uuid2502/merged_vector"

# merge_db(vector_base_folder, final_folder)----------

# query_from_vector("summary about networks", final_folder)----------------

# print(read_document(r"data/uuid2502/apple.com.filtered.json"))------------

# create_vector(read_document(r"data/uuid2502/apple.com.filtered.json"), vector_base_folder)--------------

